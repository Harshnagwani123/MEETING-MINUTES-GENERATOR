from flask import Flask, request, render_template, send_file
from faster_whisper import WhisperModel
from transformers import pipeline
from textblob import TextBlob
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import datetime
import os

app = Flask(__name__)

# Load summarization model
summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")

# Sentiment analysis function
def analyze_sentiment(text):
    blob = TextBlob(text)
    sentiment_score = blob.sentiment.polarity
    if sentiment_score > 0:
        sentiment = "Positive"
    elif sentiment_score < 0:
        sentiment = "Negative"
    else:
        sentiment = "Neutral"
    return sentiment, sentiment_score

# Export functions
def export_to_pdf(summary, file_name="meeting_summary.pdf"):
    doc = SimpleDocTemplate(file_name)
    styles = getSampleStyleSheet()
    content = [Paragraph(summary, styles['Normal'])]
    doc.build(content)

def export_to_docx(summary, file_name="meeting_summary.docx"):
    doc = Document()
    doc.add_heading('Meeting Summary', level=1)
    doc.add_paragraph(summary)
    doc.save(file_name)

def extract_key_points(summary, num_points=3):
    # Extract keywords using CountVectorizer
    vectorizer = CountVectorizer(max_features=50, stop_words="english")
    X = vectorizer.fit_transform([summary])
    terms = vectorizer.get_feature_names_out()
    
    # Compute importance by word frequency
    importance = X.sum(axis=0).A1
    important_terms = sorted(zip(importance, terms), reverse=True)
    
    # Select the top 'num_points' terms
    key_points = [term for _, term in important_terms[:num_points]]
    return "\n".join(f"- {key}" for key in key_points)

def extract_key_points(summary, num_points=3):
    # Extract keywords using CountVectorizer
    vectorizer = CountVectorizer(max_features=50, stop_words="english")
    X = vectorizer.fit_transform([summary])
    terms = vectorizer.get_feature_names_out()
    
    # Compute importance by word frequency
    importance = X.sum(axis=0).A1
    important_terms = sorted(zip(importance, terms), reverse=True)
    
    # Select the top 'num_points' terms
    key_points = [term for _, term in important_terms[:num_points]]
    return "\n".join(f"- {key}" for key in key_points)

def format_summarized_text(summary, meeting_title="Meeting Summary"):
    current_date = datetime.datetime.now().strftime("%Y-%m-%d")
    sentiment, score = analyze_sentiment(summary)

    # Generate Key Points
    key_points = extract_key_points(summary)

    # Split summary into sentences
    sentences = summary.split(". ")
    formatted_summary = "\n".join(f"- {sentence.strip()}." for sentence in sentences if sentence.strip())

    # Properly formatted output
    final_output = f"""
    =======================================
              {meeting_title}
    =======================================
    Date: {current_date}

    **Summary**:
    {formatted_summary}

    **Key Points**:
    {key_points}

    **Sentiment**:
    {sentiment} ({score:.2f})
    =======================================
    """
    return final_output

@app.route('/')
def home():
    return render_template('index.html')

@app.route("/transcribe", methods=["POST"])
def transcribe_audio():
    if 'file' not in request.files:
        return "No file part", 400
    file = request.files['file']
    if file.filename == '':
        return "No selected file", 400
    # Save the uploaded file
    if not os.path.exists("uploaded_audio"):
        os.makedirs("uploaded_audio")
    file_path = os.path.join("uploaded_audio", file.filename)
    file.save(file_path)
    # Transcribe audio
    model = WhisperModel("tiny")
    segments, info = model.transcribe(file_path)
    transcription = " ".join(segment.text for segment in segments)
    # Summarize transcription
    summary_level = request.form.get("summary_level", "brief")
    if summary_level == "brief":
        summary = summarizer(transcription, max_length=100, min_length=30, do_sample=False)[0]['summary_text']
    elif summary_level == "detailed":
        summary = summarizer(transcription, max_length=300, min_length=100, do_sample=False)[0]['summary_text']
    else:
        summary = transcription
    meeting_summary = format_summarized_text(summary)
    # Export to selected format
    format_option = request.form.get("format", "txt")
    file_name = os.path.splitext(file.filename)[0] + "_summary"
    if format_option == "pdf":
        export_to_pdf(meeting_summary, f"{file_name}.pdf")
        return send_file(f"{file_name}.pdf", as_attachment=True)
    elif format_option == "docx":
        export_to_docx(meeting_summary, f"{file_name}.docx")
        return send_file(f"{file_name}.docx", as_attachment=True)
    else:
        with open(f"{file_name}.txt", 'w', encoding='utf-8') as f:
            f.write(meeting_summary)
        return send_file(f"{file_name}.txt", as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)